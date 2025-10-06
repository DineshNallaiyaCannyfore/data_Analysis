import tempfile
import httpx
import duckdb
import pandas as pd
import docx
import re
import fitz
import os
from typing import  List
from fastapi import File, UploadFile
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain_community.vectorstores.duckdb import DuckDB
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document 
from bs4 import BeautifulSoup
import matplotlib.pyplot as plt
import base64
import io
import easyocr
from fastapi.responses import JSONResponse
import asyncio
from tqdm import tqdm

model = SentenceTransformer('all-MiniLM-L6-v2')
# Set Gemini key
os.environ["GOOGLE_API_KEY"] = "AIzaSyAto7z8ff4DnyXMZQHYiU9ubu0lsrgwt-o"
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")
CON = duckdb.connect(database=":memory:")
VECTOR_STORE = None

async def extractText(files: List[UploadFile]) -> str:
    global VECTOR_STORE
    query_text = ""
    file_text = ""
    for f in files:
        suffix = f.filename.split(".")[-1].lower()
        print("suffix",suffix,f.filename.split(".")[-1].lower())
        if suffix == "txt":
            file_text = (await f.read()).decode("utf-8")
            query_text += file_text

        elif suffix == "pdf":
            file_text = ""
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(await f.read())
                tmp_path = tmp.name
            doc = fitz.open(tmp_path)
            file_text = "".join([page.get_text() for page in doc])
            await embed_chunks_in_memory_duckdb(file_text, f.filename, suffix)

        elif suffix in ["csv"]:
            file_text = ""
            with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
                tmp.write(await f.read())
                tmp_path = tmp.name
            df = pd.read_csv(tmp_path)
            file_text = df.to_string()
            await embed_chunks_in_memory_duckdb(file_text, f.filename, suffix)

        elif suffix in ["xls", "xlsx"]:
            file_text = ""
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(await f.read())
                tmp_path = tmp.name
            df = pd.read_excel(tmp_path)
            file_text = df.to_string()
            await embed_chunks_in_memory_duckdb(file_text, f.filename, suffix)

        elif suffix in ["docx", "doc"]:
            file_text = ""
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(await f.read())
                tmp_path = tmp.name
            doc = docx.Document(tmp_path)
            file_text = "\n".join([para.text for para in doc.paragraphs])
            await embed_chunks_in_memory_duckdb(file_text, f.filename, suffix)

        elif suffix in ["jpg"]:
            file_text = ""
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(await f.read())
                tmp_path = tmp.name
                reader = easyocr.Reader(['en'])  # loads English model
                result = reader.readtext(tmp_path)
                texts = [res[1] for res in result]
                print("imgText:", " ".join(texts))
                await embed_chunks_in_memory_duckdb(file_text, f.filename, suffix)

            file_text = "\n".join(texts)

    return {
            "query": query_text,
            "fileTypeList": file_text
        }
    
def getURLS(text):
    url_pattern = r'https?://[^\s]+'
    urls = re.findall(url_pattern, text)
    return urls
    
async def extract_text_from_urls(url_list):
    if url_list: 
        all_text = ""
        async with httpx.AsyncClient() as client:
            for url in url_list:
                response = await client.get(url, timeout=10)
                soup = BeautifulSoup(response.text, "html.parser")
                text = " ".join([p.get_text() for p in soup.find_all(["p","th","td"])])
                all_text += text + "\n"
        print("all_text",len(all_text))
        await embed_chunks_in_memory_duckdb(all_text, 'webURLTest', 'txt')
        return all_text

async def embed_chunks_in_memory_duckdb(text: str, metadata: str, file_type: str):
    global VECTOR_STORE, CON

    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            add_start_index=True
        )
        chunks = text_splitter.split_text(text)
        print(f"Split text into {len(chunks)} chunks")
        metadata_obj = {
            "source": metadata,
            "file_type": file_type
        }
        documents = [
            Document(page_content=chunk, metadata=metadata_obj)
            for chunk in chunks
        ]
        embedding_model = HuggingFaceEmbeddings(model_name="intfloat/e5-base-v2")

        print("Embedding model loaded successfully")
        
        if VECTOR_STORE is None:
            print("Initializing VECTOR_STORE...")
            VECTOR_STORE = DuckDB(
                connection=CON,
                embedding=embedding_model,
                table_name="documents"
            )
        batch_size = 300
        for i in tqdm(range(0, len(documents), batch_size), desc="Embedding batches"):
            batch = documents[i:i + batch_size]
            await asyncio.to_thread(
                DuckDB.from_documents,
                documents=batch,
                embedding=embedding_model,
                connection=CON,
                table_name="documents"
            )

        print("VECTOR_STORE updated successfully")
        return VECTOR_STORE

    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }
        
async def vector_search(query: str, k: int):
    global VECTOR_STORE
    if isinstance(query, dict):
        query = query.get("query", "") 

    if not isinstance(query, str) or not query.strip():
        return {"error": "Query must be a non-empty string"}

    results = VECTOR_STORE.similarity_search(query, k=k)

    return [
        {
            "content": doc.page_content,
            "metadata": doc.metadata
        }
        for doc in results
    ]

async def generate_questions(query: str):
    docs = await vector_search(query,10)
    context = "\n\n".join([d["content"] for d in docs])
    print("context",len(context))
    prompt = ChatPromptTemplate.from_template("""
        You are an expert assistant.
        
        Context:
        {context}
        Question:
        {question}
        
        - Here is a sample Answer:
          If the answer is not matched just send the actual error in Answer property
            Eg:
                ```json
                 [{{
                    "question": "How many $2 bn movies were released before 2000?",
                    "Answer": "CRL MP(MD)/4399/2023 of Vinoth Vs The Inspector of Police"
                 }}]
            '''
        - If the answer is not matched, just send the actual error in the Answer property.

        - If the question involves a chart, graph, or visualization:
        - First, return the JSON array of objects as above for textual answers.
        - Then, in a **separate block**, return ONLY valid Python matplotlib code using "df" wrapped like this:
        
        ```python
        # matplotlib code here
        ```
        - Do NOT return matplotlib code as a string inside JSON.
        - Keep responses concise and based ONLY on the provided context.

    """)

    chain = prompt | llm
    response = await chain.ainvoke({
        "context": context,
        "question": query
    })
    print("-----response------",len(response.content))
    withImageRes = extract_and_generate_base64(response.content)
    print("----withImageRes----",len(withImageRes))
    return withImageRes

async def handleFileUpload(files: List[UploadFile] = File(...)) :
    if not files:
        return {
            "status": 200,
            "error": "Please add the file",
        }
    extractedText = await extractText(files) 
   
    if not extractedText.get('query'):  
        return {
            "status": 200,
            "error": "txt file missing"
        }
    urlResult = getURLS(extractedText.get('query'))
    if (len(urlResult) == 0) and (extractedText.get("fileTypeList") is None):
        return {
            "status": "200",
            "error": "Given file doesn't have data source, please attach url or other doc",
        }
    print("urlResult",urlResult)
    webtext = await extract_text_from_urls(urlResult)
    print("webtext",len(webtext))
    result = await generate_questions(extractedText)
    print("-------result-",result)
    return {
        "status": "200",
        "generated_answer": str(result)
    }
def extract_and_generate_base64(response_text):
    code_blocks = re.findall(r"```python(.*?)```", response_text, re.DOTALL)

    for code in code_blocks:
        try:
            code_clean = code.strip()
            code_clean = code_clean.encode("utf-8").decode("unicode_escape")
            code_clean = re.sub(r"plt\.show\(\)", "", code_clean)
            local_ns = {}
            exec(code_clean, {"plt": plt, "pd": pd}, local_ns)
            fig = plt.gcf()
            buffer = io.BytesIO()
            dpi = 150
            fig.savefig(buffer, format="png", dpi=dpi)
            buffer.seek(0)
            while buffer.getbuffer().nbytes > 100_000:
                dpi -= 10
                buffer = io.BytesIO()
                fig.savefig(buffer, format="png", dpi=dpi)
                buffer.seek(0)
                if dpi < 50:
                    break
            img_base64 = base64.b64encode(buffer.read()).decode("utf-8")
            plt.close(fig)
            response_text = response_text.replace(
                f"```python{code}```",
                f"![Chart](data:image/png;base64,{img_base64})"
            )

        except Exception as e:
            print("Error generating plot:", e, "\nCode was:\n", code_clean)

    return response_text
   
